"""
Day 16 - Step 0: Build real-world source documents.

Source text: "Pride and Prejudice" by Jane Austen (public domain, Project Gutenberg).
We build:
  1. native.pdf   - a 50+ page native (text-layer) PDF via reportlab
  2. document.docx - a DOCX with real heading styles (chapters) via python-docx
  3. document.txt  - the plain text file
  4. scanned.pdf  - an image-only PDF (no text layer) built by rasterizing a
                    few pages of native.pdf, simulating a scanned document
                    for the OCR pipeline.
"""

import os
import re
import textwrap

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
import docx
from docx.shared import Pt

RAW_TXT = "/home/claude/pride.txt"
DATA_DIR = "/home/claude/day16_project/data"
os.makedirs(DATA_DIR, exist_ok=True)

with open(RAW_TXT, "r", encoding="utf-8", errors="ignore") as f:
    raw = f.read()

# Strip Project Gutenberg header/footer boilerplate
start_marker = "PRIDE AND PREJUDICE"
start_idx = raw.find(start_marker, raw.find("*** START"))
if start_idx == -1:
    start_idx = raw.find("*** START")
end_idx = raw.find("*** END")
if end_idx == -1:
    end_idx = len(raw)
body = raw[start_idx:end_idx].strip()

# Split into chapters on lines like "Chapter 1", "Chapter I", etc.
chapter_pattern = re.compile(r"\n\s*(Chapter\s+\d+)\s*\n", re.IGNORECASE)
parts = chapter_pattern.split(body)

# parts[0] is preamble (title etc.), then alternating (heading, text)
chapters = []
if len(parts) > 1:
    preamble = parts[0]
    for i in range(1, len(parts), 2):
        heading = parts[i].strip()
        text = parts[i + 1].strip() if i + 1 < len(parts) else ""
        chapters.append((heading, text))
else:
    preamble = body
    chapters = [("Full Text", body)]

# Keep a healthy chunk of the book — enough chapters to comfortably exceed 50 PDF pages
chapters = chapters[:20]

print(f"Parsed {len(chapters)} chapters from source text.")

# -------------------------------------------------------------------
# 1. document.txt — plain text version
# -------------------------------------------------------------------
txt_path = os.path.join(DATA_DIR, "document.txt")
with open(txt_path, "w", encoding="utf-8") as f:
    f.write("PRIDE AND PREJUDICE\nby Jane Austen\n\n")
    for heading, text in chapters:
        f.write(f"{heading}\n\n{text}\n\n")
print(f"Wrote {txt_path}")

# -------------------------------------------------------------------
# 2. document.docx — DOCX with real heading styles
# -------------------------------------------------------------------
docx_path = os.path.join(DATA_DIR, "document.docx")
d = docx.Document()

title = d.add_heading("Pride and Prejudice", level=0)
sub = d.add_paragraph("by Jane Austen")
sub.style = d.styles["Subtitle"] if "Subtitle" in [s.name for s in d.styles] else d.styles["Normal"]

for heading, text in chapters:
    d.add_heading(heading.title(), level=1)
    # Split chapter text into paragraphs on blank lines
    for para in re.split(r"\n\s*\n", text):
        para = para.strip().replace("\n", " ")
        if para:
            p = d.add_paragraph(para)
            p.style.font.size = Pt(11)

d.save(docx_path)
print(f"Wrote {docx_path}")

# -------------------------------------------------------------------
# 3. native.pdf — 50+ page native text PDF (has an embedded text layer)
# -------------------------------------------------------------------
pdf_path = os.path.join(DATA_DIR, "native.pdf")

PAGE_W, PAGE_H = LETTER
MARGIN = 0.9 * inch
FONT_NAME = "Helvetica"
FONT_SIZE = 10.5
LEADING = 14
usable_width = PAGE_W - 2 * MARGIN

c = canvas.Canvas(pdf_path, pagesize=LETTER)


def new_page_cursor():
    return PAGE_H - MARGIN


y = new_page_cursor()


def draw_line(text, font=FONT_NAME, size=FONT_SIZE, extra_leading=0):
    global y
    if y < MARGIN:
        c.showPage()
        y = new_page_cursor()
    c.setFont(font, size)
    c.drawString(MARGIN, y, text)
    y -= (LEADING + extra_leading)


def draw_wrapped_paragraph(text, font=FONT_NAME, size=FONT_SIZE):
    global y
    lines = simpleSplit(text, font, size, usable_width)
    for ln in lines:
        draw_line(ln, font, size)
    y -= 6  # paragraph spacing


# Title page
draw_line("PRIDE AND PREJUDICE", font="Helvetica-Bold", size=20)
draw_line("by Jane Austen", font="Helvetica-Oblique", size=13)
c.showPage()
y = new_page_cursor()

for heading, text in chapters:
    draw_line(heading.upper(), font="Helvetica-Bold", size=14, extra_leading=6)
    for para in re.split(r"\n\s*\n", text):
        para = " ".join(para.strip().split())
        if para:
            draw_wrapped_paragraph(para)

c.save()

# Report page count
import PyPDF2
with open(pdf_path, "rb") as f:
    n_pages = len(PyPDF2.PdfReader(f).pages)
print(f"Wrote {pdf_path} ({n_pages} pages)")

# -------------------------------------------------------------------
# 4. scanned.pdf — image-only PDF (no text layer), for OCR testing
#    Rasterize the first 5 pages of native.pdf into images, then
#    rebuild a PDF from *only* those images (no text layer at all).
# -------------------------------------------------------------------
import pdfplumber
from PIL import Image
import io

scanned_pdf_path = os.path.join(DATA_DIR, "scanned.pdf")
N_SCAN_PAGES = 5

images = []
with pdfplumber.open(pdf_path) as pdf:
    for i in range(N_SCAN_PAGES):
        page = pdf.pages[i]
        pil_img = page.to_image(resolution=200).original  # PIL Image, no text layer
        images.append(pil_img.convert("RGB"))

images[0].save(scanned_pdf_path, save_all=True, append_images=images[1:])
print(f"Wrote {scanned_pdf_path} ({len(images)} rasterized/image-only pages)")

print("\nAll source documents generated in:", DATA_DIR)
